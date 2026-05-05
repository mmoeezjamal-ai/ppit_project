"""
Shared utilities for Phase 1 & Phase 2.
OCR, preprocessing, formatting analysis, and docx generation.
"""
import os
import re
import platform

import cv2
import numpy as np
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Tesseract path (Windows) ──────────────────────────────────────────────────
def configure_tesseract():
    if platform.system() == "Windows":
        for p in [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            os.path.join(os.environ.get("LOCALAPPDATA", ""), "Programs", "Tesseract-OCR", "tesseract.exe"),
        ]:
            if os.path.exists(p):
                pytesseract.pytesseract.tesseract_cmd = p
                return

configure_tesseract()

# ══════════════════════════════════════════════════════════════════════════════
# IMAGE PREPROCESSING
# ══════════════════════════════════════════════════════════════════════════════

def deskew(gray: np.ndarray) -> np.ndarray:
    try:
        coords = np.column_stack(np.where(gray < 128))
        if len(coords) < 20:
            return gray
        angle = cv2.minAreaRect(coords)[-1]
        angle = -(90 + angle) if angle < -45 else -angle
        if abs(angle) < 0.4:
            return gray
        h, w = gray.shape[:2]
        M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
        return cv2.warpAffine(gray, M, (w, h),
                               flags=cv2.INTER_CUBIC,
                               borderMode=cv2.BORDER_REPLICATE)
    except Exception:
        return gray


def get_skew_angle(gray: np.ndarray) -> float:
    try:
        coords = np.column_stack(np.where(gray < 128))
        if len(coords) < 20:
            return 0.0
        angle = cv2.minAreaRect(coords)[-1]
        return -(90 + angle) if angle < -45 else -angle
    except Exception:
        return 0.0


def _remove_ruled_lines(img_rgb: np.ndarray) -> np.ndarray:
    """
    Remove blue/colored ruled notebook lines using HSV masking.
    Returns a cleaned grayscale using the minimum-channel trick so that
    ALL colored ink (red, blue, green) becomes as dark as black ink.
    """
    hsv = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2HSV)
    hue, sat, val = hsv[:, :, 0], hsv[:, :, 1], hsv[:, :, 2]

    # Blue ruled-line mask: H≈100-130 in OpenCV (blue), moderate saturation
    blue_mask = ((hue > 88) & (hue < 138) & (sat > 25) & (val > 80)).astype(np.uint8)
    # Horizontal dilation — fills gaps along horizontal line strokes
    h_kern    = np.ones((1, 30), np.uint8)
    blue_mask = cv2.dilate(blue_mask, h_kern)

    # Minimum-channel trick: min(R,G,B) per pixel
    #   Any saturated color (red, green, blue) has at least one very LOW channel
    #   → min() picks that low value → colored text appears DARK (good for OCR)
    #   White bg: min(240,240,240)=240 (bright)   Black text: min(50,50,50)=50 (dark)
    #   Red text:  min(200,55,55)=55 (dark ✓)     Blue lines: we blank with mask
    gray_min = np.min(img_rgb, axis=2).astype(np.uint8)

    # Blank out detected ruled lines → replace with white
    gray_min[blue_mask > 0] = 255

    # Also erase long horizontal runs via morphology (catches non-blue faint lines)
    line_kern = cv2.getStructuringElement(cv2.MORPH_RECT, (60, 1))
    h_lines   = cv2.morphologyEx(gray_min, cv2.MORPH_OPEN, line_kern)
    # Only erase where the line is bright enough to be background
    gray_min[h_lines > 200] = 255

    return gray_min


def preprocess_image(pil_img: Image.Image, strategy: str = "enhanced") -> Image.Image:
    """
    strategy: 'basic' | 'enhanced' | 'aggressive'
    Full pipeline:
      1. Upscale to ≥1400px wide
      2. Invert if dark-background (UI screenshots, dark-mode docs)
      3. Remove colored ruled lines (notebook paper) via HSV + min-channel
      4. Denoise → CLAHE → sharpen → adaptive threshold → deskew → morph-close
    """
    w, h = pil_img.size
    if w < 1400:
        scale = 1400 / w
        pil_img = pil_img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    img_rgb = np.array(pil_img.convert("RGB"))

    # ── Step 1: Invert dark-background images (UI, dark-mode screenshots) ──────
    gray_check = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2GRAY)
    if np.mean(gray_check) < 110:
        img_rgb   = cv2.bitwise_not(img_rgb)

    # ── Step 2: Smart grayscale with colored-line removal ──────────────────────
    # Detect if image has colored ruled lines (notebook paper)
    hsv_check = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2HSV)
    sat_mean  = float(np.mean(hsv_check[:, :, 1]))
    has_color = sat_mean > 18  # colored ink or ruled lines present

    if has_color:
        gray = _remove_ruled_lines(img_rgb)
    else:
        gray = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2GRAY)

    if strategy == "basic":
        binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        return Image.fromarray(binary)

    if strategy in ("enhanced", "aggressive"):
        h_val = 10 if strategy == "enhanced" else 16
        gray  = cv2.fastNlMeansDenoising(gray, h=h_val, templateWindowSize=7,
                                          searchWindowSize=21)
        clahe = cv2.createCLAHE(clipLimit=3.5 if strategy == "enhanced" else 4.5,
                                  tileGridSize=(8, 8))
        gray  = clahe.apply(gray)

        if strategy == "aggressive":
            kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]], dtype=np.float32)
            gray   = np.clip(cv2.filter2D(gray, -1, kernel), 0, 255).astype(np.uint8)

        block  = 19 if strategy == "enhanced" else 25
        binary = cv2.adaptiveThreshold(gray, 255,
                                        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                        cv2.THRESH_BINARY, block, 8)
        binary = deskew(binary)
        k      = np.ones((2, 2), np.uint8)
        binary = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, k)
        return Image.fromarray(binary)

    return Image.fromarray(gray)


# ══════════════════════════════════════════════════════════════════════════════
# IMAGE QUALITY METRICS
# ══════════════════════════════════════════════════════════════════════════════

def analyze_image_quality(pil_img: Image.Image) -> dict:
    """Return quality metrics used by the Perception Agent."""
    img_rgb = np.array(pil_img.convert("RGB"))

    # Invert dark-background images before analysis
    gray_raw   = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2GRAY)
    is_inverted = float(np.mean(gray_raw)) < 110
    if is_inverted:
        img_rgb = cv2.bitwise_not(img_rgb)

    # Detect colored image (notebook paper, colored ink)
    hsv_check = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2HSV)
    sat_mean  = float(np.mean(hsv_check[:, :, 1]))
    has_color = sat_mean > 18
    has_ruled_lines = has_color  # same heuristic

    # Use cleaned grayscale for all metrics
    if has_color:
        gray_for_analysis = _remove_ruled_lines(img_rgb)
    else:
        gray_for_analysis = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2GRAY)

    brightness = float(np.mean(gray_for_analysis))
    contrast   = float(np.std(gray_for_analysis))

    lap_var = float(cv2.Laplacian(gray_for_analysis, cv2.CV_64F).var())
    sharpness = min(lap_var / 20.0, 100.0)

    noise = float(np.std(gray_for_analysis - cv2.GaussianBlur(gray_for_analysis, (5, 5), 0)))

    skew = abs(get_skew_angle(cv2.threshold(gray_for_analysis, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]))

    # Handwritten vs printed heuristic
    edges       = cv2.Canny(gray_for_analysis, 50, 150)
    edge_ratio  = float(np.sum(edges > 0)) / edges.size
    is_handwritten = edge_ratio < 0.08

    # Multi-column detection (simple): check for wide vertical white stripe in center
    col_slice  = gray[:, gray.shape[1]//2 - 20: gray.shape[1]//2 + 20]
    is_multi_col = float(np.mean(col_slice)) > 200

    quality_score = (
        min(brightness / 200 * 30, 30) +
        min(contrast  / 60  * 25, 25) +
        min(sharpness / 100 * 25, 25) +
        max(20 - noise, 0) / 1.0 * 0.1 * 20
    )
    quality_score = min(float(quality_score), 100.0)

    return {
        "brightness":     brightness,
        "contrast":       contrast,
        "sharpness":      sharpness,
        "noise":          noise,
        "skew_angle":     skew,
        "is_handwritten": is_handwritten,
        "is_multi_column":is_multi_col,
        "quality_score":  quality_score,
        "edge_ratio":     edge_ratio,
        "is_inverted":      is_inverted,
        "has_ruled_lines":  has_ruled_lines,
        "sat_mean":         sat_mean,
    }


def recommend_strategy(metrics: dict) -> tuple[str, str]:
    """Return (strategy, reasoning) based on quality metrics."""
    q = metrics["quality_score"]
    if q >= 65:
        return "enhanced", "Good quality image. Standard enhanced pipeline sufficient."
    if q >= 40:
        return "aggressive", (
            f"Moderate quality (score {q:.0f}/100). "
            "Applying aggressive denoising and sharpening."
        )
    return "aggressive", (
        f"Low quality (score {q:.0f}/100). "
        "Aggressive preprocessing to recover as much text as possible."
    )


# ══════════════════════════════════════════════════════════════════════════════
# OCR
# ══════════════════════════════════════════════════════════════════════════════

def run_ocr(pil_img: Image.Image, psm: int = 6) -> dict | None:
    config = f"--oem 3 --psm {psm} -c tessedit_char_blacklist=|"
    try:
        return pytesseract.image_to_data(pil_img, config=config,
                                         output_type=pytesseract.Output.DICT)
    except Exception:
        return None


def run_ocr_best(pil_img: Image.Image) -> tuple[dict | None, int, float]:
    """Try multiple PSM modes and return (best_data, best_psm, avg_conf)."""
    candidates = [6, 4, 11, 3]
    best_data, best_psm, best_conf = None, 6, 0.0
    for psm in candidates:
        data = run_ocr(pil_img, psm=psm)
        if data is None:
            continue
        words = [w.strip() for w, c in zip(data["text"], data["conf"])
                 if w.strip() and int(c) > 10]
        confs = [int(c) for c in data["conf"] if int(c) > 0]
        if not words:
            continue
        avg_c = sum(confs) / len(confs) if confs else 0
        score = len(words) * 0.5 + avg_c * 0.5  # weight both coverage and confidence
        if score > best_conf:
            best_conf = score
            best_data = data
            best_psm  = psm
    return best_data, best_psm, best_conf


# ══════════════════════════════════════════════════════════════════════════════
# FORMATTING ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

_BULLET_RE   = re.compile(
    r"^([•·●○◦▪▫\-–—\*]|\d{1,2}[.)]\s|[a-zA-Z][.)]\s|\([ivxIVX\d]+\))\s*"
)
_HEADING_CLN = re.compile(r"^#+\s*")


def _group_words_to_lines(data: dict) -> dict:
    lines: dict = {}
    for i, word in enumerate(data["text"]):
        word = word.strip()
        if not word or int(data["conf"][i]) < 10:
            continue
        key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        b = lines.setdefault(key, {"words": [], "left": [], "top": [],
                                    "width": [], "height": [], "conf": []})
        b["words"].append(word)
        b["left"].append(data["left"][i])
        b["top"].append(data["top"][i])
        b["width"].append(data["width"][i])
        b["height"].append(data["height"][i])
        b["conf"].append(int(data["conf"][i]))
    return lines


def detect_alignment(left: int, width: int, page_w: int) -> str:
    cx = left + width / 2
    if abs(cx - page_w / 2) < page_w * 0.12 and min(left, page_w - (left + width)) > page_w * 0.08:
        return "center"
    if left > page_w * 0.45 and page_w - (left + width) < page_w * 0.1:
        return "right"
    return "left"


def classify_line(text: str) -> tuple[bool, bool, str]:
    """Return (is_heading, is_bullet, clean_text)."""
    t = text.strip()
    if t.startswith("#"):
        return True, False, _HEADING_CLN.sub("", t).strip()
    if t.isupper() and 3 <= len(t.split()) <= 7:
        return True, False, t
    if _BULLET_RE.match(t):
        return False, True, _BULLET_RE.sub("", t).strip()
    return False, False, t


def extract_formatted_lines(ocr_data: dict, page_w: int) -> list[dict]:
    lines_dict = _group_words_to_lines(ocr_data)
    if not lines_dict:
        return []

    all_h    = [h for ld in lines_dict.values() for h in ld["height"] if h > 0]
    mean_h   = float(np.mean(all_h)) if all_h else 20.0
    result   = []

    for key in sorted(lines_dict.keys()):
        ld   = lines_dict[key]
        text = " ".join(ld["words"]).strip()
        if not text:
            continue

        left   = min(ld["left"])
        top    = min(ld["top"])
        right  = max(l + w for l, w in zip(ld["left"], ld["width"]))
        width  = right - left
        avg_h  = float(np.mean(ld["height"]))
        avg_c  = float(np.mean(ld["conf"]))

        is_h, is_b, clean = classify_line(text)
        align  = detect_alignment(left, width, page_w)
        bold   = is_h or (avg_h > mean_h * 1.28)
        h_lvl  = (1 if avg_h > mean_h * 1.5 else 2) if is_h else 0

        result.append({
            "text": text, "clean": clean,
            "top": top, "left": left, "width": width,
            "avg_h": avg_h, "mean_h": mean_h,
            "is_heading": is_h, "h_level": h_lvl,
            "is_bullet": is_b, "alignment": align,
            "is_bold": bold, "confidence": avg_c,
            "block": key[0], "par": key[1],
        })
    return result


def group_into_paragraphs(lines: list[dict]) -> list[list[dict]]:
    if not lines:
        return []
    paras, cur = [], [lines[0]]
    for i in range(1, len(lines)):
        prev, curr = lines[i - 1], lines[i]
        gap = curr["top"] - (prev["top"] + prev["avg_h"])
        if gap > max(prev["avg_h"], curr["avg_h"]) * 1.4 or curr["block"] != prev["block"]:
            paras.append(cur)
            cur = [curr]
        else:
            cur.append(curr)
    paras.append(cur)
    return paras


# ══════════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT GENERATION
# ══════════════════════════════════════════════════════════════════════════════

def _set_para_fmt(para, align: str, before: float = 0, after: float = 5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.alignment    = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                       "right":  WD_ALIGN_PARAGRAPH.RIGHT
                       }.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def generate_docx(all_pages: list, title: str = "Converted Document") -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    tp  = doc.add_paragraph()
    tr  = tp.add_run(title)
    tr.bold = True; tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x0D, 0x21, 0x37)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(4)
    sep = doc.add_paragraph("─" * 60)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(10)

    for pidx, page_paras in enumerate(all_pages):
        if len(all_pages) > 1:
            pl  = doc.add_paragraph()
            plr = pl.add_run(f"— Page {pidx + 1} —")
            plr.bold = True; plr.font.size = Pt(9)
            plr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pl.paragraph_format.space_after = Pt(6)

        for para_lines in page_paras:
            for line in para_lines:
                text = line["clean"].strip()
                if not text:
                    continue
                if line["is_heading"]:
                    lvl = line["h_level"] if line["h_level"] in (1, 2) else 2
                    p   = doc.add_heading(text, level=lvl)
                    _set_para_fmt(p, line["alignment"], before=8, after=3)
                elif line["is_bullet"]:
                    p   = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(text)
                    run.font.size = Pt(11)
                    if line["is_bold"]:
                        run.bold = True
                    _set_para_fmt(p, "left", before=1, after=2)
                else:
                    p   = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.size = Pt(11)
                    if line["is_bold"]:
                        run.bold = True
                    _set_para_fmt(p, line["alignment"], before=2, after=3)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)

        if pidx < len(all_pages) - 1:
            doc.add_page_break()
    return doc
