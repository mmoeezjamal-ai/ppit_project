import streamlit as st
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import cv2
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import os
import platform

# ── Tesseract path (Windows local support) ───────────────────────────────────
def _configure_tesseract():
    if platform.system() == "Windows":
        candidates = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            os.path.join(os.environ.get("LOCALAPPDATA", ""), "Programs", "Tesseract-OCR", "tesseract.exe"),
        ]
        for p in candidates:
            if os.path.exists(p):
                pytesseract.pytesseract.tesseract_cmd = p
                return
_configure_tesseract()

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Image → Word Converter | FAST-NUCES",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Global CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');

html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

.hero {
    background: linear-gradient(135deg, #0d2137 0%, #1a4a7a 60%, #2563a8 100%);
    color: white;
    padding: 2rem 2.5rem;
    border-radius: 14px;
    margin-bottom: 1.5rem;
    text-align: center;
    box-shadow: 0 4px 24px rgba(13,33,55,0.18);
}
.hero h1 { font-size: 2.2rem; font-weight: 700; margin: 0; letter-spacing: -0.5px; }
.hero p  { font-size: 1rem; margin: 6px 0 0; opacity: 0.88; }

.step-card {
    background: #f0f6ff;
    border-left: 4px solid #2563a8;
    padding: 10px 14px;
    border-radius: 6px;
    margin: 6px 0;
    font-size: 0.92rem;
}

.stat-banner {
    display: flex; gap: 12px; flex-wrap: wrap; margin-top: 10px;
}
.stat-pill {
    background: #e8f0fe;
    border: 1px solid #c5d8f8;
    border-radius: 20px;
    padding: 5px 14px;
    font-size: 0.85rem;
    color: #1a4a7a;
    font-weight: 600;
}

.section-title {
    font-size: 1.15rem;
    font-weight: 700;
    color: #0d2137;
    border-bottom: 2px solid #2563a8;
    padding-bottom: 5px;
    margin-bottom: 10px;
}

.success-card {
    background: #d1fae5;
    border: 1px solid #6ee7b7;
    border-radius: 10px;
    padding: 18px 22px;
    color: #065f46;
    font-weight: 600;
    font-size: 1.05rem;
    text-align: center;
}

.warning-card {
    background: #fef3c7;
    border: 1px solid #fcd34d;
    border-radius: 10px;
    padding: 14px 18px;
    color: #92400e;
    font-size: 0.9rem;
}

.tag { display:inline-block; background:#dbeafe; color:#1e40af;
       border-radius:4px; padding:2px 8px; font-size:0.78rem; margin:2px; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# IMAGE PREPROCESSING
# ══════════════════════════════════════════════════════════════════════════════

def _deskew(gray: np.ndarray) -> np.ndarray:
    """Detect and correct skew angle of a binarised image."""
    try:
        coords = np.column_stack(np.where(gray < 128))
        if len(coords) < 20:
            return gray
        angle = cv2.minAreaRect(coords)[-1]
        angle = -(90 + angle) if angle < -45 else -angle
        if abs(angle) < 0.4:
            return gray
        h, w = gray.shape[:2]
        M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
        return cv2.warpAffine(gray, M, (w, h),
                               flags=cv2.INTER_CUBIC,
                               borderMode=cv2.BORDER_REPLICATE)
    except Exception:
        return gray


def preprocess_image(pil_img: Image.Image, enhance: bool = True) -> Image.Image:
    """
    Full preprocessing pipeline:
      1. Upscale if too small
      2. Greyscale
      3. Denoise
      4. CLAHE contrast
      5. Adaptive threshold
      6. Deskew
    """
    # Upscale to at least 1400 px wide so Tesseract sees large glyphs
    w, h = pil_img.size
    if w < 1400:
        scale = 1400 / w
        pil_img = pil_img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    img = np.array(pil_img.convert("RGB"))
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

    if not enhance:
        return Image.fromarray(gray)

    # Denoise
    gray = cv2.fastNlMeansDenoising(gray, h=12, templateWindowSize=7, searchWindowSize=21)

    # CLAHE – adaptive contrast
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)

    # Sharpen slightly
    kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]], dtype=np.float32)
    gray = cv2.filter2D(gray, -1, kernel)
    gray = np.clip(gray, 0, 255).astype(np.uint8)

    # Adaptive binarisation
    binary = cv2.adaptiveThreshold(
        gray, 255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        19, 9,
    )

    # Deskew
    binary = _deskew(binary)

    # Morphological close – fill small gaps in ink
    k = np.ones((2, 2), np.uint8)
    binary = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, k)

    return Image.fromarray(binary)


# ══════════════════════════════════════════════════════════════════════════════
# OCR ENGINE
# ══════════════════════════════════════════════════════════════════════════════

def run_ocr(pil_img: Image.Image) -> dict | None:
    """Run Tesseract and return word-level data dict."""
    config = r"--oem 3 --psm 6 -c tessedit_char_blacklist=|"
    try:
        return pytesseract.image_to_data(
            pil_img,
            config=config,
            output_type=pytesseract.Output.DICT,
        )
    except Exception as e:
        st.error(f"OCR failed: {e}")
        return None


def _group_words_to_lines(data: dict) -> dict:
    """Aggregate OCR word tokens into logical lines keyed by (block, par, line)."""
    lines: dict = {}
    for i, word in enumerate(data["text"]):
        word = word.strip()
        if not word or int(data["conf"][i]) < 10:
            continue
        key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        bucket = lines.setdefault(key, {"words": [], "left": [], "top": [],
                                        "width": [], "height": [], "conf": []})
        bucket["words"].append(word)
        bucket["left"].append(data["left"][i])
        bucket["top"].append(data["top"][i])
        bucket["width"].append(data["width"][i])
        bucket["height"].append(data["height"][i])
        bucket["conf"].append(int(data["conf"][i]))
    return lines


# ══════════════════════════════════════════════════════════════════════════════
# FORMATTING ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

_BULLET_RE = re.compile(
    r"^([•·●○◦▪▫\-–—\*]|\d{1,2}[.)]\s|[a-zA-Z][.)]\s|\([ivxIVX\d]+\))\s*"
)
_HEADING_CLEAN_RE = re.compile(r"^#+\s*")


def _detect_alignment(left: int, width: int, page_w: int) -> str:
    cx = left + width / 2
    mar_l = left
    mar_r = page_w - (left + width)
    if abs(cx - page_w / 2) < page_w * 0.12 and min(mar_l, mar_r) > page_w * 0.08:
        return "center"
    if mar_l > page_w * 0.45 and mar_r < page_w * 0.1:
        return "right"
    return "left"


def _classify_line(text: str):
    """Return (is_heading, is_bullet, clean_text)."""
    t = text.strip()
    # Heading: starts with one or more # chars
    if t.startswith("#"):
        return True, False, _HEADING_CLEAN_RE.sub("", t).strip()
    # Heading: ALL CAPS short phrase
    if t.isupper() and 3 <= len(t.split()) <= 7:
        return True, False, t
    # Bullet
    if _BULLET_RE.match(t):
        clean = _BULLET_RE.sub("", t).strip()
        return False, True, clean
    return False, False, t


def extract_formatted_lines(ocr_data: dict, page_w: int) -> list[dict]:
    """Convert raw OCR token dict into a list of annotated line records."""
    lines_dict = _group_words_to_lines(ocr_data)
    if not lines_dict:
        return []

    all_h = [h for ld in lines_dict.values() for h in ld["height"] if h > 0]
    mean_h = float(np.mean(all_h)) if all_h else 20.0

    result = []
    for key in sorted(lines_dict.keys()):
        ld = lines_dict[key]
        text = " ".join(ld["words"]).strip()
        if not text:
            continue

        left   = min(ld["left"])
        top    = min(ld["top"])
        right  = max(l + w for l, w in zip(ld["left"], ld["width"]))
        width  = right - left
        avg_h  = float(np.mean(ld["height"]))
        avg_c  = float(np.mean(ld["conf"]))

        is_heading, is_bullet, clean = _classify_line(text)
        alignment  = _detect_alignment(left, width, page_w)
        is_bold    = is_heading or (avg_h > mean_h * 1.28)
        # Estimate heading level by relative font size
        if is_heading:
            h_level = 1 if avg_h > mean_h * 1.5 else 2
        else:
            h_level = 0

        result.append({
            "text":      text,
            "clean":     clean,
            "top":       top,
            "left":      left,
            "width":     width,
            "avg_h":     avg_h,
            "is_heading": is_heading,
            "h_level":   h_level,
            "is_bullet": is_bullet,
            "alignment": alignment,
            "is_bold":   is_bold,
            "confidence": avg_c,
            "block":     key[0],
            "par":       key[1],
        })
    return result


def group_into_paragraphs(lines: list[dict]) -> list[list[dict]]:
    """Cluster lines into paragraphs using vertical gap heuristic."""
    if not lines:
        return []
    paras, cur = [], [lines[0]]
    for i in range(1, len(lines)):
        prev, curr = lines[i - 1], lines[i]
        gap = curr["top"] - (prev["top"] + prev["avg_h"])
        # New paragraph: big vertical gap OR new block
        if gap > max(prev["avg_h"], curr["avg_h"]) * 1.4 or curr["block"] != prev["block"]:
            paras.append(cur)
            cur = [curr]
        else:
            cur.append(curr)
    paras.append(cur)
    return paras


# ══════════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT GENERATION
# ══════════════════════════════════════════════════════════════════════════════

def _set_para_fmt(para, align: str, before: float = 0, after: float = 5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.alignment    = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def generate_docx(all_paras: list[list[list[dict]]], title: str = "Converted Document") -> Document:
    """
    Build a python-docx Document from structured paragraph/line data.
    all_paras is a list of pages; each page is a list of paragraphs;
    each paragraph is a list of line dicts.
    """
    doc = Document()

    # ── Margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Cover title
    tp = doc.add_paragraph()
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x0D, 0x21, 0x37)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(4)

    sep = doc.add_paragraph("─" * 60)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(10)

    for page_idx, page_paras in enumerate(all_paras):
        # Page label
        if len(all_paras) > 1:
            pl = doc.add_paragraph()
            plr = pl.add_run(f"— Page {page_idx + 1} —")
            plr.bold = True
            plr.font.size = Pt(9)
            plr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pl.paragraph_format.space_after = Pt(6)

        for para_lines in page_paras:
            if not para_lines:
                continue

            for line in para_lines:
                text = line["clean"].strip()
                if not text:
                    continue

                if line["is_heading"]:
                    level = line["h_level"] if line["h_level"] in (1, 2) else 2
                    p = doc.add_heading(text, level=level)
                    _set_para_fmt(p, line["alignment"], before=8, after=3)

                elif line["is_bullet"]:
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(text)
                    run.font.size = Pt(11)
                    if line["is_bold"]:
                        run.bold = True
                    _set_para_fmt(p, "left", before=1, after=2)

                else:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.size = Pt(11)
                    if line["is_bold"]:
                        run.bold = True
                    _set_para_fmt(p, line["alignment"], before=2, after=3)

            # Blank line between paragraphs
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)

        # Page break between pages (not after last)
        if page_idx < len(all_paras) - 1:
            doc.add_page_break()

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# STREAMLIT UI
# ══════════════════════════════════════════════════════════════════════════════

def main():
    # ── Hero banner
    st.markdown("""
    <div class="hero">
        <h1>📄 Image → Word Converter</h1>
        <p>FAST-NUCES &nbsp;|&nbsp; BSAI &nbsp;|&nbsp; PPIT Project Phase 1</p>
        <p style="font-size:0.88rem;opacity:0.75;">
            Tesseract OCR &nbsp;·&nbsp; Formatting Detection &nbsp;·&nbsp; .docx Export
        </p>
    </div>
    """, unsafe_allow_html=True)

    # ── Sidebar
    with st.sidebar:
        st.markdown("### ⚙️ Options")
        enhance   = st.checkbox("🔧 Enhance image quality", value=True,
                                help="Denoising, CLAHE, adaptive threshold, deskew")
        show_pre  = st.checkbox("👁️ Show preprocessed image", value=False)
        show_raw  = st.checkbox("📋 Show raw OCR text", value=True)
        doc_title = st.text_input("📝 Document title", value="Converted Document")

        st.markdown("---")
        st.markdown("### 📊 Detected Formatting")
        st.markdown("""
        <span class="tag">📰 H1 / H2 Headings</span>
        <span class="tag">📌 Bullet lists</span>
        <span class="tag">↔️ L / C / R align</span>
        <span class="tag">🔤 Bold text</span>
        <span class="tag">📄 Paragraphs</span>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### ℹ️ About")
        st.markdown("""
        **Stack:** Python · Streamlit · Tesseract · OpenCV · python-docx  
        **Input:** JPG · JPEG · PNG · BMP · TIFF  
        **Output:** Formatted `.docx`  
        """)
        st.markdown("---")
        st.caption("FAST-NUCES BSAI — PPIT Project · 2026")

    # ── Upload
    st.markdown('<div class="section-title">📁 Upload Images</div>', unsafe_allow_html=True)
    uploaded = st.file_uploader(
        "Drag & drop scanned pages or handwritten notes (JPG / PNG)",
        type=["jpg", "jpeg", "png", "bmp", "tiff"],
        accept_multiple_files=True,
    )

    if not uploaded:
        c1, c2, c3, c4 = st.columns(4)
        c1.info("📰 Heading detection\n\n`#` prefix or ALL CAPS")
        c2.info("📌 Bullet detection\n\n`•` `-` `1.` patterns")
        c3.info("↔️ Alignment\n\nLeft / Center / Right")
        c4.info("🔤 Bold text\n\nFont size heuristic")
        st.markdown("""
        <div class="warning-card">
        ⚠️ <b>Note:</b> Tesseract OCR works best on <b>printed / typed</b> text.
        For handwritten documents it will do its best — image enhancement is
        applied automatically to improve recognition.
        </div>
        """, unsafe_allow_html=True)
        return

    # ── Convert button
    if not st.button("🚀 Convert to Word Document", type="primary", use_container_width=True):
        st.info(f"✅ {len(uploaded)} image(s) ready. Press **Convert** to start.")
        return

    # ── Processing ────────────────────────────────────────────────────────────
    progress = st.progress(0, text="Starting…")
    all_pages: list[list[list[dict]]] = []   # pages > paragraphs > lines
    total_stats = {"paras": 0, "headings": 0, "bullets": 0, "bold": 0}

    for idx, f in enumerate(uploaded):
        step_pct = idx / len(uploaded)
        progress.progress(step_pct, text=f"Processing {f.name} …")

        pil_img = Image.open(f).convert("RGB")
        page_w  = pil_img.width

        with st.expander(f"📷  Page {idx + 1} — {f.name}", expanded=True):
            col_img, col_txt = st.columns([1, 1])

            with col_img:
                st.markdown("**Original image**")
                st.image(pil_img, use_container_width=True)

            # Preprocess
            preprocessed = preprocess_image(pil_img, enhance=enhance)

            if show_pre:
                with col_txt:
                    st.markdown("**Preprocessed (fed to OCR)**")
                    st.image(preprocessed, use_container_width=True)

            # OCR
            with st.spinner("Running Tesseract OCR…"):
                ocr_data = run_ocr(preprocessed)

            if ocr_data is None:
                st.error("OCR failed — skipping this image.")
                all_pages.append([])
                continue

            # Analyse
            fmt_lines = extract_formatted_lines(ocr_data, page_w)
            page_paras = group_into_paragraphs(fmt_lines)
            all_pages.append(page_paras)

            # Stats
            h_cnt = sum(1 for p in page_paras for l in p if l["is_heading"])
            b_cnt = sum(1 for p in page_paras for l in p if l["is_bullet"])
            bo_cnt= sum(1 for p in page_paras for l in p if l["is_bold"])
            total_stats["paras"]    += len(page_paras)
            total_stats["headings"] += h_cnt
            total_stats["bullets"]  += b_cnt
            total_stats["bold"]     += bo_cnt

            # Show extracted text
            if not show_pre:
                raw_text = "\n".join(
                    " ".join(l["text"] for l in para) for para in page_paras
                )
                with col_txt:
                    if show_raw:
                        st.markdown("**Extracted text**")
                        st.text_area("", raw_text, height=400, key=f"raw_{idx}", label_visibility="collapsed")
                    else:
                        st.markdown("**Formatting breakdown**")
                        for para in page_paras:
                            for line in para:
                                tags = []
                                if line["is_heading"]:
                                    tags.append(f"H{line['h_level']}")
                                if line["is_bullet"]:
                                    tags.append("bullet")
                                if line["is_bold"]:
                                    tags.append("bold")
                                tags.append(line["alignment"])
                                st.markdown(
                                    f"`{'·'.join(tags)}` {line['clean'][:80]}…"
                                    if len(line["clean"]) > 80 else
                                    f"`{'·'.join(tags)}` {line['clean']}"
                                )

            # Per-page metrics
            m1, m2, m3, m4 = st.columns(4)
            m1.metric("Paragraphs",  len(page_paras))
            m2.metric("Headings",    h_cnt)
            m3.metric("Bullets",     b_cnt)
            m4.metric("Bold lines",  bo_cnt)

    progress.progress(0.92, text="Generating Word document…")

    # ── Generate docx ─────────────────────────────────────────────────────────
    try:
        doc = generate_docx(all_pages, title=doc_title)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        progress.progress(1.0, text="Done!")

        st.markdown("""
        <div class="success-card">✅ Word document generated successfully!</div>
        """, unsafe_allow_html=True)

        # Overall stats
        st.markdown("### 📊 Overall Summary")
        s1, s2, s3, s4, s5 = st.columns(5)
        s1.metric("Pages",      len(uploaded))
        s2.metric("Paragraphs", total_stats["paras"])
        s3.metric("Headings",   total_stats["headings"])
        s4.metric("Bullets",    total_stats["bullets"])
        s5.metric("Bold lines", total_stats["bold"])

        st.download_button(
            label="⬇️  Download  " + doc_title + ".docx",
            data=buf,
            file_name=f"{doc_title.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
        )
        st.balloons()

    except Exception as exc:
        st.error(f"Document generation error: {exc}")
        st.exception(exc)


if __name__ == "__main__":
    main()
